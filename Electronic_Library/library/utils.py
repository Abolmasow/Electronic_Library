import json
import csv
import io
import os
from datetime import datetime
from decimal import Decimal
from django.http import HttpResponse
from django.db.models import QuerySet
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import openpyxl
from openpyxl.styles import Font, Alignment
from docx import Document

class DataExporter:
    """Class for exporting data in different formats."""
    
    @staticmethod
    def export_to_json(queryset: QuerySet, fields: list) -> HttpResponse:
        """Export data to JSON format."""
        data = list(queryset.values(*fields))
        
        response = HttpResponse(
            json.dumps(data, ensure_ascii=False, indent=2, default=str),
            content_type='application/json; charset=utf-8'
        )
        response['Content-Disposition'] = 'attachment; filename="export.json"'
        
        return response
    
    @staticmethod
    def export_to_csv(queryset: QuerySet, fields: list, field_names: list) -> HttpResponse:
        """Export data to CSV format."""
        response = HttpResponse(
            content_type='text/csv; charset=utf-8'
        )
        response['Content-Disposition'] = 'attachment; filename="export.csv"'
        
        writer = csv.writer(response)
        writer.writerow(field_names)
        
        for item in queryset:
            row = []
            for field in fields:
                value = getattr(item, field, '')
                
                if hasattr(value, 'all'):  # Many-to-many field
                    value = ', '.join(str(v) for v in value.all())
                elif isinstance(value, Decimal):
                    value = str(value)
                elif value is None:
                    value = ''
                
                row.append(str(value))
            
            writer.writerow(row)
        
        return response
    
    @staticmethod
    def export_to_pdf(queryset: QuerySet, title: str, columns: list) -> HttpResponse:
        """Export data to PDF format."""
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{title}.pdf"'
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        
        elements = []
        styles = getSampleStyleSheet()
        
        elements.append(Paragraph(title, styles['Title']))
        elements.append(Paragraph(f"Дата формирования: {datetime.now().strftime('%Y-%m-%d %H:%M')}", 
                                 styles['Normal']))
        
        data = [columns]
        
        for item in queryset:
            row = []
            for column in columns:
                attr_name = column.lower().replace(' ', '_')
                value = getattr(item, attr_name, '')
                
                if hasattr(value, 'all'):
                    value = ', '.join(str(v) for v in value.all()[:3])
                elif isinstance(value, Decimal):
                    value = f"{value:.2f}"
                elif value is None:
                    value = ''
                
                row.append(str(value))
            data.append(row)
        
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        elements.append(table)
        doc.build(elements)
        
        pdf = buffer.getvalue()
        buffer.close()
        response.write(pdf)
        
        return response
    
    @staticmethod
    def export_to_docx(queryset: QuerySet, title: str, columns: list) -> HttpResponse:
        """Export data to DOCX format."""
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{title}.docx"'
        
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(f"Дата формирования: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = 'Table Grid'
        
        header_cells = table.rows[0].cells
        for i, column in enumerate(columns):
            header_cells[i].text = column
        
        for item in queryset:
            row_cells = table.add_row().cells
            for i, column in enumerate(columns):
                attr_name = column.lower().replace(' ', '_')
                value = getattr(item, attr_name, '')
                row_cells[i].text = str(value)
        
        doc.save(response)
        return response
    
    @staticmethod
    def export_to_xlsx(queryset: QuerySet, title: str, columns: list) -> HttpResponse:
        """Export data to XLSX format."""
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="{title}.xlsx"'
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = title[:31]
        
        ws['A1'] = title
        ws['A1'].font = Font(size=16, bold=True)
        ws.merge_cells('A1:E1')
        
        ws['A2'] = f"Дата формирования: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        
        for col_num, column in enumerate(columns, 1):
            cell = ws.cell(row=4, column=col_num)
            cell.value = column
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal='center')
        
        for row_num, item in enumerate(queryset, 5):
            for col_num, column in enumerate(columns, 1):
                cell = ws.cell(row=row_num, column=col_num)
                attr_name = column.lower().replace(' ', '_')
                value = getattr(item, attr_name, '')
                cell.value = value
        
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        wb.save(response)
        return response